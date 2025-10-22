# process_doc_utils.py
import os
import time
import xml.etree.ElementTree as ET
from typing import Dict, Any, Optional, List
from xml.dom import minidom

# Optional: DOCX
try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    Document = None

# Optional: Diagrams (Graphviz required)
try:
    from diagrams import Diagram, Cluster, Edge
    from diagrams.generic.blank import Blank
    DIAGRAMS_AVAILABLE = True
except Exception:
    DIAGRAMS_AVAILABLE = False

# Optional: BPMN Python library
try:
    from bpmn_python.bpmn_diagram_rep import BpmnDiagramGraph
    from bpmn_python import bpmn_diagram_visualizer
    BPMN_PYTHON_AVAILABLE = True
except ImportError:
    BPMN_PYTHON_AVAILABLE = False

# Configuration
ENABLE_BPMN_GENERATION = os.getenv("ENABLE_BPMN_GENERATION", "false").lower() == "true"
ENABLE_PNG_GENERATION = os.getenv("ENABLE_PNG_GENERATION", "false").lower() == "true"
USE_BPMN_PYTHON = os.getenv("USE_BPMN_PYTHON", "false").lower() == "true"

# Utility functions
def slug(s: str) -> str:
    return "".join(c.lower() if c.isalnum() else "_" for c in s)[:40]

def generate_bpmn_xml(model: Dict[str, Any]) -> str:
    """
    Enhanced BPMN 2.0 XML generator with Car-Wash.bpmn inspired features:
    - Separate processes for each actor 
    - Exclusive gateways for decision points
    - Message flows for inter-process communication
    - Proper task flow references (incoming/outgoing)
    - Valid sequence flow connections with conditional branches
    - Compliant with BPMN 2.0 schema validation
    """
    if not ENABLE_BPMN_GENERATION:
        return ""

    # Use the exact namespace pattern from the valid bpmn.io example
    BPMN_NS = "http://www.omg.org/spec/BPMN/20100524/MODEL"
    
    # Create definitions matching the valid bpmn.io example structure exactly
    defs = ET.Element("bpmn2:definitions")
    defs.set("xmlns:bpmn2", BPMN_NS)
    defs.set("xmlns:bpmndi", "http://www.omg.org/spec/BPMN/20100524/DI")
    defs.set("xmlns:dc", "http://www.omg.org/spec/DD/20100524/DC")
    defs.set("xmlns:di", "http://www.omg.org/spec/DD/20100524/DI")
    defs.set("id", "empty-definitions")
    defs.set("targetNamespace", "http://bpmn.io/schema/bpmn")

    def b(tag): return f"bpmn2:{tag}"
    
    # Define namespace map for XML operations
    namespaces = {"bpmn2": BPMN_NS}

    # Extract actors from the model
    actors = model.get("actors", [])
    main_flow_steps = model.get("main_flow", [])
    
    # Get unique actors from main flow if not defined in actors list
    if not actors and main_flow_steps:
        unique_actors = set()
        for step in main_flow_steps:
            actor = step.get("actor", "").strip()
            if actor:
                unique_actors.add(actor)
        actors = [{"name": actor, "role": actor, "responsibilities": ""} for actor in unique_actors]
    
    # If still no actors, create a default one
    if not actors:
        actors = [{"name": "Process Owner", "role": "Process Owner", "responsibilities": "Executes the process"}]

    process_id = f"Process_{slug(model.get('title') or 'process')}"
    
    # Create collaboration with separate processes for each actor
    collaboration = ET.SubElement(defs, b("collaboration"))
    collaboration.set("id", "Collaboration_1")
    
    # Create participants and process mappings
    actor_processes = {}
    for i, actor in enumerate(actors, 1):
        actor_name = actor.get("name", f"Actor {i}")
        process_id = f"Process_{slug(actor_name)}_{i}"
        actor_processes[actor_name] = process_id
        
        participant = ET.SubElement(collaboration, b("participant"))
        participant.set("id", f"Participant_{i}")
        participant.set("name", actor_name[:50])
        participant.set("processRef", process_id)
    
    # Group tasks by actor
    tasks_by_actor = {}
    for step in main_flow_steps:
        actor = step.get("actor", "").strip()
        if not actor and actors:
            actor = actors[0].get("name", "Process Owner")
        if actor not in tasks_by_actor:
            tasks_by_actor[actor] = []
        tasks_by_actor[actor].append(step)
    
    # Create separate processes for each actor
    process_elements = {}
    base_process_name = model.get('title', 'Business Process')
    for actor_name, process_id in actor_processes.items():
        process = ET.SubElement(defs, b("process"))
        process.set("id", process_id)
        process.set("isExecutable", "false")
        process.set("name", f"{base_process_name} - {actor_name}")
        process_elements[actor_name] = process

    # Documentation element for overall process description
    if model.get('overview') and process_elements:
        # Add documentation to the first process
        first_process = list(process_elements.values())[0]
        doc = ET.SubElement(first_process, b("documentation"))
        doc.text = model.get('overview')[:300]

    # Create start and end events for each process
    start_events = {}
    end_events = {}
    for actor_name, process in process_elements.items():
        # Start event
        start = ET.SubElement(process, b("startEvent"))
        start.set("id", f"StartEvent_{slug(actor_name)}")
        start_events[actor_name] = start.get("id")
        
        # End event  
        end = ET.SubElement(process, b("endEvent"))
        end.set("id", f"EndEvent_{slug(actor_name)}")
        end_events[actor_name] = end.get("id")
    
    # Create tasks, gateways and flows for each actor in their respective processes
    task_elements = {}
    gateway_elements = {}
    sequence_flows = []
    message_flows = []
    
    # Build main flow tasks grouped by actor with enhanced logic
    for actor_name, actor_tasks in tasks_by_actor.items():
        process = process_elements.get(actor_name)
        if not process:
            continue
            
        prev_element_id = start_events[actor_name]
        
        # Detect decision points and escalation patterns
        has_decision_point = False
        decision_keywords = ["determine", "decide", "analyze", "evaluate", "choose", "review", "check"]
        escalation_keywords = ["escalate", "supervisor", "technical", "policy"]
        assignment_keywords = ["assign", "forward", "route", "send", "transfer", "delegate"]
        
        for i, step in enumerate(actor_tasks, start=1):
            task_id = f"Task_{slug(actor_name)}_{i}"
            label = step.get("action") or step.get("id") or f"Step {i}"
            
            # Check if this step involves decision making, escalation, or assignment
            is_decision = any(keyword in label.lower() for keyword in decision_keywords)
            has_escalation = any(keyword in label.lower() for keyword in escalation_keywords)
            is_assignment = any(keyword in label.lower() for keyword in assignment_keywords)
            
            # Create task in the actor's process (using generic 'task' element like Car-Wash.bpmn)
            task = ET.SubElement(process, b("task"))
            task.set("id", task_id)
            task.set("name", label[:80])
            task_elements[task_id] = task
            
            # Add sequence flow from previous element to this task
            sequence_flows.append({
                "id": f"Flow_{slug(actor_name)}_{i}",
                "sourceRef": prev_element_id,
                "targetRef": task_id,
                "process": process
            })
            
            # If this is an assignment task, add gateway with message flows to other participants
            if is_assignment:
                gateway_id = f"Gateway_{slug(actor_name)}_{i}"
                gateway = ET.SubElement(process, b("exclusiveGateway"))
                gateway.set("id", gateway_id)
                gateway.set("name", f"Assignment Decision: {label[:30]}?")
                gateway_elements[gateway_id] = gateway
                
                # Flow from task to gateway
                sequence_flows.append({
                    "id": f"Flow_{slug(actor_name)}_{i}_to_gateway",
                    "sourceRef": task_id,
                    "targetRef": gateway_id,
                    "process": process
                })
                
                # Create assignment flows to other participants
                if "technical" in label.lower() and "Technical Support Team" in [a.get("name", "") for a in actors]:
                    # Assignment to Technical Support
                    sequence_flows.append({
                        "id": f"Flow_{slug(actor_name)}_{i}_to_technical",
                        "sourceRef": gateway_id,
                        "targetRef": end_events[actor_name],
                        "process": process,
                        "name": "To Technical Support"
                    })
                    # Message flow to technical support
                    technical_task_id = "Task_technical_support_team_1"
                    if technical_task_id in task_elements:
                        message_flow = ET.SubElement(collaboration, b("messageFlow"))
                        message_flow.set("id", f"MessageFlow_{slug(actor_name)}_to_technical")
                        message_flow.set("name", "Assignment")
                        message_flow.set("sourceRef", task_id)
                        message_flow.set("targetRef", technical_task_id)
                        message_flows.append(message_flow)
                
                if "policy" in label.lower() and "Policy Review Team" in [a.get("name", "") for a in actors]:
                    # Assignment to Policy Review
                    sequence_flows.append({
                        "id": f"Flow_{slug(actor_name)}_{i}_to_policy",
                        "sourceRef": gateway_id,
                        "targetRef": end_events[actor_name],
                        "process": process,
                        "name": "To Policy Review"
                    })
                    # Message flow to policy review
                    policy_task_id = "Task_policy_review_team_1"
                    if policy_task_id in task_elements:
                        message_flow = ET.SubElement(collaboration, b("messageFlow"))
                        message_flow.set("id", f"MessageFlow_{slug(actor_name)}_to_policy")
                        message_flow.set("name", "Assignment")
                        message_flow.set("sourceRef", task_id)
                        message_flow.set("targetRef", policy_task_id)
                        message_flows.append(message_flow)
                
                prev_element_id = gateway_id
                has_decision_point = True
            
            # If this is a decision point, add an exclusive gateway after the task
            elif is_decision and i < len(actor_tasks):
                gateway_id = f"Gateway_{slug(actor_name)}_{i}"
                gateway = ET.SubElement(process, b("exclusiveGateway"))
                gateway.set("id", gateway_id)
                gateway.set("name", f"Decision: {label[:30]}?")
                gateway_elements[gateway_id] = gateway
                
                # Flow from task to gateway
                sequence_flows.append({
                    "id": f"Flow_{slug(actor_name)}_{i}_to_gateway",
                    "sourceRef": task_id,
                    "targetRef": gateway_id,
                    "process": process
                })
                
                # Create conditional flows from gateway with different targets
                if has_escalation:
                    # Simple path - continue to next task
                    if i < len(actor_tasks):
                        next_task_id = f"Task_{slug(actor_name)}_{i+1}"
                        sequence_flows.append({
                            "id": f"Flow_{slug(actor_name)}_{i}_simple",
                            "sourceRef": gateway_id,
                            "targetRef": next_task_id,
                            "process": process,
                            "name": "Simple Case"
                        })
                    else:
                        # If last task, go to end
                        sequence_flows.append({
                            "id": f"Flow_{slug(actor_name)}_{i}_simple",
                            "sourceRef": gateway_id,
                            "targetRef": end_events[actor_name],
                            "process": process,
                            "name": "Simple Case"
                        })
                    
                    # Escalation path - create escalation task with message flow
                    escalation_task_id = f"Task_{slug(actor_name)}_{i}_escalation"
                    escalation_task = ET.SubElement(process, b("task"))
                    escalation_task.set("id", escalation_task_id)
                    escalation_task.set("name", f"Escalate: {label[:50]}")
                    task_elements[escalation_task_id] = escalation_task
                    
                    sequence_flows.append({
                        "id": f"Flow_{slug(actor_name)}_{i}_escalate",
                        "sourceRef": gateway_id,
                        "targetRef": escalation_task_id,
                        "process": process,
                        "name": "Complex Case"
                    })
                    
                    # Add message flow for escalation to supervisor
                    if "supervisor" in label.lower() and "Supervisor" in [a.get("name", "") for a in actors]:
                        supervisor_task_id = "Task_supervisor_1"
                        if supervisor_task_id in task_elements:
                            message_flow = ET.SubElement(collaboration, b("messageFlow"))
                            message_flow.set("id", f"MessageFlow_{slug(actor_name)}_escalation")
                            message_flow.set("name", "Escalation")
                            message_flow.set("sourceRef", escalation_task_id)
                            message_flow.set("targetRef", supervisor_task_id)
                            message_flows.append(message_flow)
                    
                    # Connect escalation task to end or next task
                    if i < len(actor_tasks):
                        next_task_id = f"Task_{slug(actor_name)}_{i+1}"
                        sequence_flows.append({
                            "id": f"Flow_{slug(actor_name)}_{i}_escalation_continue",
                            "sourceRef": escalation_task_id,
                            "targetRef": next_task_id,
                            "process": process
                        })
                    else:
                        sequence_flows.append({
                            "id": f"Flow_{slug(actor_name)}_{i}_escalation_end",
                            "sourceRef": escalation_task_id,
                            "targetRef": end_events[actor_name],
                            "process": process
                        })
                else:
                    # Non-escalation decision gateway - simple yes/no paths
                    if i < len(actor_tasks):
                        next_task_id = f"Task_{slug(actor_name)}_{i+1}"
                        # Yes path
                        sequence_flows.append({
                            "id": f"Flow_{slug(actor_name)}_{i}_yes",
                            "sourceRef": gateway_id,
                            "targetRef": next_task_id,
                            "process": process,
                            "name": "Yes"
                        })
                        # No path - skip to end
                        sequence_flows.append({
                            "id": f"Flow_{slug(actor_name)}_{i}_no",
                            "sourceRef": gateway_id,
                            "targetRef": end_events[actor_name],
                            "process": process,
                            "name": "No"
                        })
                    else:
                        # Last task decision - only one path to end
                        sequence_flows.append({
                            "id": f"Flow_{slug(actor_name)}_{i}_complete",
                            "sourceRef": gateway_id,
                            "targetRef": end_events[actor_name],
                            "process": process,
                            "name": "Complete"
                        })
                
                prev_element_id = gateway_id
                has_decision_point = True
            else:
                prev_element_id = task_id
        
        # Connect last element to end event
        if prev_element_id != start_events[actor_name]:
            sequence_flows.append({
                "id": f"Flow_{slug(actor_name)}_end",
                "sourceRef": prev_element_id,
                "targetRef": end_events[actor_name],
                "process": process
            })
    
    # Create sequence flows within each process 
    for flow_data in sequence_flows:
        flow = ET.SubElement(flow_data["process"], b("sequenceFlow"))
        flow.set("id", flow_data["id"])
        flow.set("sourceRef", flow_data["sourceRef"])
        flow.set("targetRef", flow_data["targetRef"])
        if flow_data.get("name"):
            flow.set("name", flow_data["name"])

    # Create intelligent message flows based on process interactions
    message_flow_counter = 1
    
    # Look for escalation and communication patterns
    escalation_actors = ["Supervisor", "Technical Support Team", "Policy Review Team"]
    customer_service_actors = ["Customer Service Representative", "Customer Service", "Representative"]
    
    for source_actor in actor_processes.keys():
        for target_actor in actor_processes.keys():
            if source_actor != target_actor:
                source_tasks = [task_id for task_id in task_elements.keys() if slug(source_actor) in task_id]
                target_tasks = [task_id for task_id in task_elements.keys() if slug(target_actor) in task_id]
                
                if source_tasks and target_tasks:
                    # Create message flows for escalation patterns
                    should_create_flow = False
                    
                    # Customer service to supervisor escalation
                    if (any(cs in source_actor for cs in customer_service_actors) and 
                        "Supervisor" in target_actor):
                        escalation_task = None
                        for task_id in source_tasks:
                            if task_id in task_elements:
                                task_name = task_elements[task_id].get("name", "")
                                if "escalate" in task_name.lower():
                                    escalation_task = task_id
                                    break
                        
                        if escalation_task:
                            message_flow = ET.SubElement(collaboration, b("messageFlow"))
                            message_flow.set("id", f"MessageFlow_{message_flow_counter}")
                            message_flow.set("name", "Escalation")
                            message_flow.set("sourceRef", escalation_task)
                            message_flow.set("targetRef", target_tasks[0])
                            message_flow_counter += 1
                            should_create_flow = True
                    
                    # Supervisor to technical/policy teams
                    elif ("Supervisor" in source_actor and 
                          any(team in target_actor for team in ["Technical", "Policy"])):
                        message_flow = ET.SubElement(collaboration, b("messageFlow"))
                        message_flow.set("id", f"MessageFlow_{message_flow_counter}")
                        message_flow.set("name", "Assignment")
                        message_flow.set("sourceRef", source_tasks[-1])  # Last task of supervisor
                        message_flow.set("targetRef", target_tasks[0])   # First task of target team
                        message_flow_counter += 1
                        should_create_flow = True
                    
                    # Back to customer service for follow-up
                    elif (any(team in source_actor for team in ["Technical", "Policy", "Supervisor"]) and
                          any(cs in target_actor for cs in customer_service_actors)):
                        follow_up_task = None
                        for task_id in target_tasks:
                            if task_id in task_elements:
                                task_name = task_elements[task_id].get("name", "")
                                if "follow" in task_name.lower() or "communicate" in task_name.lower():
                                    follow_up_task = task_id
                                    break
                        
                        if follow_up_task:
                            message_flow = ET.SubElement(collaboration, b("messageFlow"))
                            message_flow.set("id", f"MessageFlow_{message_flow_counter}")
                            message_flow.set("name", "Resolution")
                            message_flow.set("sourceRef", source_tasks[-1])
                            message_flow.set("targetRef", follow_up_task)
                            message_flow_counter += 1
                            should_create_flow = True
                    
                    # Limit to 3 message flows maximum to avoid clutter
                    if should_create_flow and message_flow_counter > 4:
                        break
        
        if message_flow_counter > 4:
            break

    # Add BPMN Diagram Interchange (DI) for visual layout like Car-Wash.bpmn
    diagram = ET.SubElement(defs, "bpmndi:BPMNDiagram")
    diagram.set("id", "Collaboration_1_di")
    
    plane = ET.SubElement(diagram, "bpmndi:BPMNPlane")
    plane.set("id", "Collaboration_1_plane")
    plane.set("bpmnElement", "Collaboration_1")
    
    # Create shapes for participants (horizontal pools)
    pool_height = 250
    pool_width = 1000
    y_offset = 0
    
    for i, (actor_name, process_id) in enumerate(actor_processes.items()):
        # Participant pool shape
        participant_shape = ET.SubElement(plane, "bpmndi:BPMNShape")
        participant_shape.set("id", f"Participant_{i+1}_shape")
        participant_shape.set("bpmnElement", f"Participant_{i+1}")
        participant_shape.set("isHorizontal", "true")
        
        bounds = ET.SubElement(participant_shape, "dc:Bounds")
        bounds.set("x", "20")
        bounds.set("y", str(y_offset))
        bounds.set("width", str(pool_width))
        bounds.set("height", str(pool_height))
        
        y_offset += pool_height + 30  # Space between pools
    
    # Create shapes for tasks, events, and gateways
    task_width = 100
    task_height = 80
    event_size = 36
    gateway_size = 50
    
    x_position = 100
    for i, (actor_name, process_id) in enumerate(actor_processes.items()):
        y_base = i * (pool_height + 30) + 50  # Position within the pool
        current_x = x_position
        
        # Start event shape
        start_shape = ET.SubElement(plane, "bpmndi:BPMNShape")
        start_shape.set("id", f"StartEvent_{slug(actor_name)}_shape")
        start_shape.set("bpmnElement", f"StartEvent_{slug(actor_name)}")
        
        start_bounds = ET.SubElement(start_shape, "dc:Bounds")
        start_bounds.set("x", str(current_x))
        start_bounds.set("y", str(y_base))
        start_bounds.set("width", str(event_size))
        start_bounds.set("height", str(event_size))
        current_x += 120
        
        # Task shapes for this actor
        if actor_name in tasks_by_actor:
            for j, step in enumerate(tasks_by_actor[actor_name], 1):
                task_id = f"Task_{slug(actor_name)}_{j}"
                
                # Task shape
                task_shape = ET.SubElement(plane, "bpmndi:BPMNShape")
                task_shape.set("id", f"{task_id}_shape")
                task_shape.set("bpmnElement", task_id)
                
                task_bounds = ET.SubElement(task_shape, "dc:Bounds")
                task_bounds.set("x", str(current_x))
                task_bounds.set("y", str(y_base - 10))
                task_bounds.set("width", str(task_width))
                task_bounds.set("height", str(task_height))
                current_x += 150
                
                # Gateway shape if this task has a decision point
                gateway_id = f"Gateway_{slug(actor_name)}_{j}"
                if gateway_id in gateway_elements:
                    gateway_shape = ET.SubElement(plane, "bpmndi:BPMNShape")
                    gateway_shape.set("id", f"{gateway_id}_shape")
                    gateway_shape.set("bpmnElement", gateway_id)
                    
                    gateway_bounds = ET.SubElement(gateway_shape, "dc:Bounds")
                    gateway_bounds.set("x", str(current_x))
                    gateway_bounds.set("y", str(y_base + 10))
                    gateway_bounds.set("width", str(gateway_size))
                    gateway_bounds.set("height", str(gateway_size))
                    
                    # Gateway label
                    gateway_label = ET.SubElement(gateway_shape, "bpmndi:BPMNLabel")
                    label_bounds = ET.SubElement(gateway_label, "dc:Bounds")
                    label_bounds.set("x", str(current_x - 50))
                    label_bounds.set("y", str(y_base - 30))
                    label_bounds.set("width", str(gateway_size + 100))
                    label_bounds.set("height", "40")
                    
                    current_x += 120
                    
                    # Escalation task shape if exists
                    escalation_task_id = f"Task_{slug(actor_name)}_{j}_escalation"
                    if escalation_task_id in task_elements:
                        escalation_shape = ET.SubElement(plane, "bpmndi:BPMNShape")
                        escalation_shape.set("id", f"{escalation_task_id}_shape")
                        escalation_shape.set("bpmnElement", escalation_task_id)
                        
                        escalation_bounds = ET.SubElement(escalation_shape, "dc:Bounds")
                        escalation_bounds.set("x", str(current_x))
                        escalation_bounds.set("y", str(y_base + 60))  # Below main flow
                        escalation_bounds.set("width", str(task_width))
                        escalation_bounds.set("height", str(task_height))
                        current_x += 150
        
        # End event shape
        end_shape = ET.SubElement(plane, "bpmndi:BPMNShape")
        end_shape.set("id", f"EndEvent_{slug(actor_name)}_shape")
        end_shape.set("bpmnElement", f"EndEvent_{slug(actor_name)}")
        
        end_bounds = ET.SubElement(end_shape, "dc:Bounds")
        end_bounds.set("x", str(current_x))
        end_bounds.set("y", str(y_base))
        end_bounds.set("width", str(event_size))
        end_bounds.set("height", str(event_size))
    
    # Create edges for sequence flows with proper waypoint calculations
    shape_positions = {}
    
    # Build a map of element positions for waypoint calculations
    for i, (actor_name, process_id) in enumerate(actor_processes.items()):
        y_base = i * (pool_height + 30) + 50 + (pool_height // 2)  # Center of pool vertically
        current_x = 100 + event_size//2  # Start with event center
        
        # Start event position (center of the event)
        shape_positions[f"StartEvent_{slug(actor_name)}"] = (current_x, y_base)
        current_x += 120
        
        # Task positions for this actor with proper spacing
        if actor_name in tasks_by_actor:
            for j, step in enumerate(tasks_by_actor[actor_name], 1):
                task_id = f"Task_{slug(actor_name)}_{j}"
                # Task center position
                shape_positions[task_id] = (current_x + task_width//2, y_base)
                current_x += task_width + 50  # Better spacing
                
                # Gateway position if exists - place after task
                gateway_id = f"Gateway_{slug(actor_name)}_{j}"
                if gateway_id in gateway_elements:
                    shape_positions[gateway_id] = (current_x + gateway_size//2, y_base)
                    current_x += gateway_size + 50
                    
                    # Escalation task position if exists
                    escalation_task_id = f"Task_{slug(actor_name)}_{j}_escalation"
                    if escalation_task_id in task_elements:
                        shape_positions[escalation_task_id] = (current_x + task_width//2, y_base + 80)  # Below main flow
        
        # End event position (center of the event)
        shape_positions[f"EndEvent_{slug(actor_name)}"] = (current_x + event_size//2, y_base)
    
    # Add incoming/outgoing references to ALL elements including start and end events
    all_bpmn_elements = {**task_elements, **gateway_elements}
    
    # Add start and end events to the elements for incoming/outgoing processing
    for actor_name in process_elements.keys():
        start_id = f"StartEvent_{slug(actor_name)}"
        end_id = f"EndEvent_{slug(actor_name)}"
        
        # Find the start and end elements in the process
        process = process_elements[actor_name]
        for child in process:
            if child.get("id") == start_id:
                all_bpmn_elements[start_id] = child
            elif child.get("id") == end_id:
                all_bpmn_elements[end_id] = child
    
    # Add incoming/outgoing references to all elements
    for element_id, element in all_bpmn_elements.items():
        incoming_flows = [f["id"] for f in sequence_flows if f["targetRef"] == element_id]
        outgoing_flows = [f["id"] for f in sequence_flows if f["sourceRef"] == element_id]
        
        for flow_id in incoming_flows:
            incoming = ET.SubElement(element, b("incoming"))
            incoming.text = flow_id
            
        for flow_id in outgoing_flows:
            outgoing = ET.SubElement(element, b("outgoing"))
            outgoing.text = flow_id

    # Create sequence flow edges with CORRECTED waypoint calculations
    for flow_data in sequence_flows:
        edge = ET.SubElement(plane, "bpmndi:BPMNEdge")
        edge.set("id", f"{flow_data['id']}_edge")
        edge.set("bpmnElement", flow_data["id"])
        
        # Get source and target positions
        source_ref = flow_data["sourceRef"]
        target_ref = flow_data["targetRef"]
        
        source_pos = shape_positions.get(source_ref)
        target_pos = shape_positions.get(target_ref)
        
        if not source_pos or not target_pos:
            print(f"Warning: Missing position for {source_ref} -> {target_ref}")
            continue
        
        source_x, source_y = source_pos
        target_x, target_y = target_pos
        
        # Create waypoints for proper connection
        waypoint1 = ET.SubElement(edge, "di:waypoint")
        waypoint1.set("x", str(int(source_x)))
        waypoint1.set("y", str(int(source_y)))
        
        # Only add intermediate waypoints if there's significant vertical offset
        if abs(source_y - target_y) > 50:  # Different lanes
            # Add intermediate waypoint for better routing
            waypoint_mid = ET.SubElement(edge, "di:waypoint")
            waypoint_mid.set("x", str(int((source_x + target_x) // 2)))
            waypoint_mid.set("y", str(int(source_y)))
            
            waypoint_mid2 = ET.SubElement(edge, "di:waypoint")
            waypoint_mid2.set("x", str(int((source_x + target_x) // 2)))
            waypoint_mid2.set("y", str(int(target_y)))
        
        waypoint2 = ET.SubElement(edge, "di:waypoint")
        waypoint2.set("x", str(int(target_x)))
        waypoint2.set("y", str(int(target_y)))
        
        # Add label if the flow has a name
        if flow_data.get("name"):
            edge_label = ET.SubElement(edge, "bpmndi:BPMNLabel")
            label_bounds = ET.SubElement(edge_label, "dc:Bounds")
            mid_x = (source_x + target_x) // 2
            mid_y = (source_y + target_y) // 2 - 10
            label_bounds.set("x", str(int(mid_x - 30)))
            label_bounds.set("y", str(int(mid_y)))
            label_bounds.set("width", "60")
            label_bounds.set("height", "20")
    
    # Create edges for message flows with proper cross-pool connections
    try:
        # Find message flows using the same naming convention as creation
        message_flows = collaboration.findall("bpmn2:messageFlow", {"bpmn2": BPMN_NS})
        
        for msg_flow_elem in message_flows:
            edge = ET.SubElement(plane, "bpmndi:BPMNEdge")
            edge.set("id", f"{msg_flow_elem.get('id')}_edge")
            edge.set("bpmnElement", msg_flow_elem.get("id"))
            
            # Get source and target element positions for message flows
            source_ref = msg_flow_elem.get("sourceRef")
            target_ref = msg_flow_elem.get("targetRef")
            
            source_pos = shape_positions.get(source_ref, (500, 150))
            target_pos = shape_positions.get(target_ref, (500, 350))
            
            # Convert to actual center positions for message flows
            source_center_x = source_pos[0] + 50  # Add half of task width
            source_center_y = source_pos[1] + 40  # Add half of task height
            target_center_x = target_pos[0] + 50
            target_center_y = target_pos[1] + 40
            
            # Message flow waypoints (cross-pool connections)
            waypoint1 = ET.SubElement(edge, "di:waypoint")
            waypoint1.set("x", str(source_center_x))
            waypoint1.set("y", str(source_center_y))
            
            # Add intermediate waypoint for visual clarity if crossing pools
            if abs(source_center_y - target_center_y) > 100:  # Different pools
                waypoint_mid = ET.SubElement(edge, "di:waypoint")
                waypoint_mid.set("x", str(source_center_x))
                waypoint_mid.set("y", str((source_center_y + target_center_y) // 2))
                
                waypoint_mid2 = ET.SubElement(edge, "di:waypoint")
                waypoint_mid2.set("x", str(target_center_x))
                waypoint_mid2.set("y", str((source_center_y + target_center_y) // 2))
            
            waypoint2 = ET.SubElement(edge, "di:waypoint")
            waypoint2.set("x", str(target_center_x))
            waypoint2.set("y", str(target_center_y))
            
            # Message flow label
            if msg_flow_elem.get("name"):
                edge_label = ET.SubElement(edge, "bpmndi:BPMNLabel")
                label_bounds = ET.SubElement(edge_label, "dc:Bounds")
                mid_x = (source_center_x + target_center_x) // 2
                mid_y = (source_center_y + target_center_y) // 2
                label_bounds.set("x", str(mid_x - 40))
                label_bounds.set("y", str(mid_y - 10))
                label_bounds.set("width", "80")
                label_bounds.set("height", "20")
    except Exception as e:
        print(f"Warning: Error creating message flow edges: {e}")

    # Format XML with proper indentation and add XML declaration
    rough = ET.tostring(defs, encoding="utf-8")
    pretty = minidom.parseString(rough).toprettyxml(indent="  ", encoding="utf-8")
    
    # Convert to string and ensure proper XML declaration
    pretty_str = pretty.decode("utf-8")
    if not pretty_str.startswith('<?xml'):
        pretty_str = '<?xml version="1.0" encoding="utf-8"?>\n' + pretty_str
    
    return pretty_str


def generate_bpmn_python(model: Dict[str, Any]) -> str:
    """
    Alternative BPMN 2.0 XML generator using the bpmn-python library.
    This provides a more programmatic approach to BPMN diagram creation.
    
    Features:
    - Uses bpmn-python library for proper BPMN structure
    - Automatic layout and positioning
    - Simplified API for creating BPMN elements
    - Built-in validation and compliance
    """
    if not ENABLE_BPMN_GENERATION:
        return ""
    
    if not BPMN_PYTHON_AVAILABLE:
        print("[BPMN-Python] Library not available; falling back to manual XML generation")
        return generate_bpmn_xml(model)
    
    try:
        # Create a new BPMN diagram
        bpmn_graph = BpmnDiagramGraph()
        
        # Set diagram properties
        process_name = model.get('title', 'Business Process')
        bpmn_graph.diagram_attributes["name"] = process_name
        
        # Extract actors and main flow
        actors = model.get("actors", [])
        main_flow_steps = model.get("main_flow", [])
        
        # Get unique actors from main flow if not defined in actors list
        if not actors and main_flow_steps:
            unique_actors = set()
            for step in main_flow_steps:
                actor = step.get("actor", "").strip()
                if actor:
                    unique_actors.add(actor)
            actors = [{"name": actor, "role": actor, "responsibilities": ""} for actor in unique_actors]
        
        # If still no actors, create a default one
        if not actors:
            actors = [{"name": "Process Owner", "role": "Process Owner", "responsibilities": "Executes the process"}]
        
        # Create pools for each actor
        pools = {}
        processes = {}
        
        for i, actor in enumerate(actors, 1):
            actor_name = actor.get("name", f"Actor {i}")
            
            # Create process for this actor
            process_id = f"Process_{slug(actor_name)}_{i}"
            process_node = bpmn_graph.add_process_to_diagram(process_id)
            processes[actor_name] = process_node
            
            # Create pool (participant) for this actor
            pool_id = f"Pool_{slug(actor_name)}_{i}"
            pool_node = bpmn_graph.add_participant_to_diagram(pool_id, actor_name, process_id)
            pools[actor_name] = pool_node
        
        # Group tasks by actor
        tasks_by_actor = {}
        for step in main_flow_steps:
            actor = step.get("actor", "").strip()
            if not actor and actors:
                actor = actors[0].get("name", "Process Owner")
            if actor not in tasks_by_actor:
                tasks_by_actor[actor] = []
            tasks_by_actor[actor].append(step)
        
        # Create BPMN elements for each actor's process
        task_nodes = {}
        start_events = {}
        end_events = {}
        
        for actor_name, process_node in processes.items():
            # Add start event
            start_event_id = f"StartEvent_{slug(actor_name)}"
            start_event = bpmn_graph.add_start_event_to_diagram(process_node, start_event_id, 
                                                               start_event_name=f"Start {actor_name}")
            start_events[actor_name] = start_event
            
            # Add end event
            end_event_id = f"EndEvent_{slug(actor_name)}"
            end_event = bpmn_graph.add_end_event_to_diagram(process_node, end_event_id,
                                                           end_event_name=f"End {actor_name}")
            end_events[actor_name] = end_event
            
            # Add tasks for this actor
            if actor_name in tasks_by_actor:
                prev_node = start_event
                
                for i, step in enumerate(tasks_by_actor[actor_name], 1):
                    task_id = f"Task_{slug(actor_name)}_{i}"
                    task_name = step.get("action") or step.get("id") or f"Step {i}"
                    
                    # Determine task type based on content
                    task_name_lower = task_name.lower()
                    if any(keyword in task_name_lower for keyword in ["user", "manual", "customer"]):
                        # User task
                        task_node = bpmn_graph.add_user_task_to_diagram(process_node, task_id, task_name)
                    elif any(keyword in task_name_lower for keyword in ["service", "system", "automatic"]):
                        # Service task
                        task_node = bpmn_graph.add_service_task_to_diagram(process_node, task_id, task_name)
                    elif any(keyword in task_name_lower for keyword in ["script", "calculate", "compute"]):
                        # Script task
                        task_node = bpmn_graph.add_script_task_to_diagram(process_node, task_id, task_name)
                    else:
                        # Generic task
                        task_node = bpmn_graph.add_task_to_diagram(process_node, task_id, task_name)
                    
                    task_nodes[task_id] = task_node
                    
                    # Add sequence flow from previous node to this task
                    flow_id = f"Flow_{slug(actor_name)}_{i}"
                    bpmn_graph.add_sequence_flow_to_diagram(process_node, prev_node, task_node, 
                                                           sequence_flow_id=flow_id)
                    
                    # Check for decision points and add gateways
                    decision_keywords = ["determine", "decide", "analyze", "evaluate", "choose", "review", "check"]
                    if any(keyword in task_name_lower for keyword in decision_keywords):
                        # Add exclusive gateway after decision task
                        gateway_id = f"Gateway_{slug(actor_name)}_{i}"
                        gateway_name = f"Decision: {task_name[:30]}?"
                        gateway_node = bpmn_graph.add_exclusive_gateway_to_diagram(process_node, gateway_id, 
                                                                                 gateway_name=gateway_name)
                        
                        # Flow from task to gateway
                        gateway_flow_id = f"Flow_{slug(actor_name)}_{i}_to_gateway"
                        bpmn_graph.add_sequence_flow_to_diagram(process_node, task_node, gateway_node,
                                                               sequence_flow_id=gateway_flow_id)
                        
                        prev_node = gateway_node
                        
                        # Add conditional flows from gateway
                        if i < len(tasks_by_actor[actor_name]):
                            # Continue to next task (Yes path)
                            next_task_id = f"Task_{slug(actor_name)}_{i+1}"
                            # We'll connect this when we create the next task
                        else:
                            # Last task - go to end event
                            end_flow_id = f"Flow_{slug(actor_name)}_{i}_to_end"
                            bpmn_graph.add_sequence_flow_to_diagram(process_node, gateway_node, end_event,
                                                                   sequence_flow_id=end_flow_id,
                                                                   sequence_flow_name="Complete")
                    else:
                        prev_node = task_node
                
                # Connect last task/gateway to end event
                if prev_node != start_event and prev_node != end_event:
                    final_flow_id = f"Flow_{slug(actor_name)}_final"
                    bpmn_graph.add_sequence_flow_to_diagram(process_node, prev_node, end_event,
                                                           sequence_flow_id=final_flow_id)
        
        # Add message flows between processes for collaboration
        message_flow_counter = 1
        escalation_actors = ["Supervisor", "Technical Support Team", "Policy Review Team"]
        customer_service_actors = ["Customer Service Representative", "Customer Service", "Representative"]
        
        for source_actor in processes.keys():
            for target_actor in processes.keys():
                if source_actor != target_actor and message_flow_counter <= 3:  # Limit message flows
                    # Create message flows for escalation patterns
                    should_create_flow = False
                    source_task = None
                    target_task = None
                    
                    # Find appropriate tasks for message flow
                    source_tasks = [task_id for task_id in task_nodes.keys() if slug(source_actor) in task_id]
                    target_tasks = [task_id for task_id in task_nodes.keys() if slug(target_actor) in task_id]
                    
                    if source_tasks and target_tasks:
                        # Customer service to supervisor escalation
                        if (any(cs in source_actor for cs in customer_service_actors) and 
                            "Supervisor" in target_actor):
                            # Find escalation task
                            for task_id in source_tasks:
                                if task_id in task_nodes:
                                    # Check if this task involves escalation
                                    if "escalate" in task_id.lower():
                                        source_task = task_nodes[task_id]
                                        target_task = task_nodes[target_tasks[0]]
                                        should_create_flow = True
                                        break
                        
                        # Supervisor to technical/policy teams
                        elif ("Supervisor" in source_actor and 
                              any(team in target_actor for team in ["Technical", "Policy"])):
                            if source_tasks and target_tasks:
                                source_task = task_nodes[source_tasks[-1]]  # Last task of supervisor
                                target_task = task_nodes[target_tasks[0]]   # First task of target team
                                should_create_flow = True
                    
                    if should_create_flow and source_task and target_task:
                        try:
                            message_flow_id = f"MessageFlow_{message_flow_counter}"
                            message_flow_name = "Escalation" if "Supervisor" in target_actor else "Assignment"
                            bpmn_graph.add_message_flow_to_diagram(source_task, target_task,
                                                                 message_flow_id=message_flow_id,
                                                                 message_flow_name=message_flow_name)
                            message_flow_counter += 1
                        except Exception as e:
                            print(f"Warning: Could not create message flow: {e}")
        
        # Add documentation if available
        if model.get('overview'):
            # Add documentation to the diagram
            try:
                # bpmn-python might not support documentation directly, so we'll add it as an annotation
                doc_text = model.get('overview')[:300]  # Limit documentation length
                bpmn_graph.diagram_attributes["documentation"] = doc_text
            except Exception:
                pass  # Documentation not supported
        
        # Export to BPMN XML
        bpmn_xml = bpmn_graph.export_xml_file_no_di()
        
        # Add XML declaration if not present
        if not bpmn_xml.startswith('<?xml'):
            bpmn_xml = '<?xml version="1.0" encoding="UTF-8"?>\n' + bpmn_xml
        
        return bpmn_xml
        
    except Exception as e:
        print(f"[BPMN-Python] Error generating BPMN: {e}")
        print("[BPMN-Python] Falling back to manual XML generation")
        return generate_bpmn_xml(model)


def generate_bpmn(model: Dict[str, Any]) -> str:
    """
    Main BPMN generation function that selects the appropriate approach based on environment variable.
    
    Environment variables:
    - USE_BPMN_PYTHON: If 'true', uses bpmn-python library approach
    - If 'false' or not set, uses manual XML generation approach
    
    Args:
        model: Dictionary containing the process model data
        
    Returns:
        str: BPMN 2.0 XML string
    """
    use_bpmn_python = os.getenv('USE_BPMN_PYTHON', 'false').lower() == 'true'
    
    if use_bpmn_python:
        print("[BPMN] Using bpmn-python library approach")
        return generate_bpmn_python(model)
    else:
        print("[BPMN] Using manual XML generation approach")
        return generate_bpmn_xml(model)


def render_diagrams_png(model: Dict[str, Any], path: str) -> Optional[str]:
    """
    Renders an improved PNG using the 'diagrams' library.
    - Better text wrapping and sizing
    - Improved layout and spacing
    - Better visual representation of process flow
    """
    if not ENABLE_PNG_GENERATION:
        print("[PNG] Generation disabled by ENABLE_PNG_GENERATION environment variable")
        return None
        
    if not DIAGRAMS_AVAILABLE:
        print("[Diagrams] Library not available or Graphviz missing; skipping PNG.")
        return None

    # Build per-actor nodes
    steps = model.get("main_flow", [])
    if not steps:
        print("[Diagrams] No main_flow; skipping PNG.")
        return None

    def wrap_text(text: str, max_length: int = 20) -> str:
        """Wrap text for better display in diagram nodes."""
        if not text or len(text) <= max_length:
            return text
        
        words = text.split()
        lines = []
        current_line = []
        current_length = 0
        
        for word in words:
            if current_length + len(word) + 1 <= max_length:
                current_line.append(word)
                current_length += len(word) + 1
            else:
                if current_line:
                    lines.append(' '.join(current_line))
                current_line = [word]
                current_length = len(word)
        
        if current_line:
            lines.append(' '.join(current_line))
        
        return '\\n'.join(lines[:3])  # Limit to 3 lines max

    def create_step_label(step: Dict[str, Any]) -> str:
        """Create a well-formatted label for a step."""
        step_id = (step.get('id') or '').strip()
        action = (step.get('action') or '').strip()
        
        if step_id and action:
            # If both ID and action, format as "ID: Action"
            full_text = f"{step_id}: {action}"
        elif action:
            # If only action, use that
            full_text = action
        elif step_id:
            # If only ID, use that
            full_text = step_id
        else:
            # Fallback
            full_text = "Process Step"
        
        return wrap_text(full_text, 25)

    # Group steps by actor
    by_actor: Dict[str, List[Dict[str, Any]]] = {}
    for s in steps:
        actor = (s.get("actor") or "Unassigned").strip() or "Unassigned"
        # Wrap actor names too
        actor = wrap_text(actor, 15)
        by_actor.setdefault(actor, []).append(s)

    graph_name = wrap_text((model.get("title") or "Process").strip() or "Process", 30)
    
    # Diagrams saves to a base name without extension
    base_name = os.path.splitext(path)[0]
    
    # Use top-down direction for better process flow visualization
    with Diagram(graph_name, filename=base_name, show=False, outformat="png", 
                 direction="TB", graph_attr={"splines": "ortho", "nodesep": "1.0", "ranksep": "1.5"}):
        
        lane_nodes: Dict[str, List[Any]] = {}
        
        # Try to import more specific diagram elements
        try:
            from diagrams.generic.compute import Rack
            from diagrams.generic.network import Router
            step_node_class = Rack
        except ImportError:
            step_node_class = Blank
        
        for actor, actor_steps in by_actor.items():
            with Cluster(actor):
                nodes = []
                for s in actor_steps:
                    label = create_step_label(s)
                    node = step_node_class(label)
                    nodes.append(node)
                lane_nodes[actor] = nodes

        # Connect main flow in order
        prev_node = None
        flat_nodes = []
        
        # Create a map from step dict to node for ordering
        step_to_node = {}
        for actor, nodes in lane_nodes.items():
            for s, n in zip(by_actor[actor], nodes):
                step_to_node[id(s)] = n
        
        # Build flat list in original step order
        for s in steps:
            n = step_to_node.get(id(s))
            if n:
                flat_nodes.append(n)
        
        # Connect nodes with improved edges
        for i, n in enumerate(flat_nodes):
            if prev_node:
                prev_node >> Edge(style="solid", color="blue") >> n
            prev_node = n

        # Add alternate paths with better positioning
        alts = model.get("alternate_paths", [])
        if alts and flat_nodes:
            for i, ap in enumerate(alts[:2]):  # Limit to 2 alternate paths
                alt_name = wrap_text(ap.get("name", f"Alt {i+1}"), 15)
                alt_node = step_node_class(alt_name)
                flat_nodes[0] >> Edge(style="dashed", color="orange") >> alt_node

        # Add exceptions with better positioning
        excs = model.get("exceptions", [])
        if excs and flat_nodes:
            for i, ex in enumerate(excs[:2]):  # Limit to 2 exceptions
                exc_name = wrap_text(ex.get("name", f"Exception {i+1}"), 15)
                exc_node = step_node_class(exc_name)
                flat_nodes[-1] >> Edge(style="dotted", color="red") >> exc_node

    return path


def render_markdown_document(md_text: str, path: str) -> None:
    """Render markdown document to file."""
    with open(path, "w", encoding="utf-8") as f:
        f.write(md_text)


def render_docx_from_markdownish(md_text: str, path: str) -> None:
    """Render DOCX document from markdown-like text."""
    if Document is None:
        print("[DOCX] python-docx not installed; skipping DOCX.")
        return
    doc = Document()
    try:
        styles = doc.styles
        normal = styles['Normal']
        normal.font.name = 'Calibri'
        normal.font.size = Pt(11)
    except Exception:
        pass

    lines = md_text.splitlines()
    in_code = False
    code_buffer = []
    for line in lines:
        if line.strip().startswith("```"):
            if in_code and code_buffer:
                # End code block
                code_text = "\n".join(code_buffer)
                p = doc.add_paragraph(code_text)
                p.style = 'No Spacing'
                try:
                    font = p.runs[0].font
                    font.name = 'Courier New'
                    font.size = Pt(10)
                except Exception:
                    pass
                code_buffer = []
            in_code = not in_code
            continue
        if in_code:
            code_buffer.append(line)
        else:
            if line.startswith("#"):
                level = len(line) - len(line.lstrip("#"))
                text = line.lstrip("# ").strip()
                if level == 1:
                    p = doc.add_heading(text, level=1)
                elif level == 2:
                    p = doc.add_heading(text, level=2)
                else:
                    p = doc.add_heading(text, level=3)
            else:
                if line.strip():
                    doc.add_paragraph(line)

    doc.save(path)